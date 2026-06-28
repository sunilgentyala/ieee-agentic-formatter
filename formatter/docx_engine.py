"""python-docx engine that produces a strictly IEEE-compliant two-column .docx."""

import math
import re
from io import BytesIO
from typing import List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from agent.models import Author, IEEEPaper

# ---------------------------------------------------------------------------
# Ordinal labels
# ---------------------------------------------------------------------------

_ORDINALS = ["1st", "2nd", "3rd", "4th", "5th", "6th", "7th", "8th", "9th", "10th"]


def _ordinal(n: int) -> str:
    return _ORDINALS[n - 1] if n <= len(_ORDINALS) else f"{n}th"


# ---------------------------------------------------------------------------
# Roman / letter helpers
# ---------------------------------------------------------------------------

_ROMAN = [
    (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
    (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
    (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
]


def _to_roman(num: int) -> str:
    result = ""
    for value, numeral in _ROMAN:
        while num >= value:
            result += numeral
            num -= value
    return result


def _to_letter(num: int) -> str:
    return chr(ord("A") + num - 1)


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _mk(tag: str) -> OxmlElement:
    return OxmlElement(tag)


def _set_font(run, name: str, size_pt: float, bold: bool = False, italic: bool = False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic


def _set_para_spacing(para, before_pt: float = 0, after_pt: float = 0, line: int = 240):
    pPr = para._p.get_or_add_pPr()
    for s in pPr.findall(qn("w:spacing")):
        pPr.remove(s)
    sp = _mk("w:spacing")
    sp.set(qn("w:before"), str(int(before_pt * 20)))
    sp.set(qn("w:after"), str(int(after_pt * 20)))
    sp.set(qn("w:line"), str(line))
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _set_page_margins(section, top=0.75, bottom=1.0, left=0.625, right=0.625):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def _set_page_size_letter(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)


# ---------------------------------------------------------------------------
# Author table helpers
# ---------------------------------------------------------------------------

def _no_borders_on_table(tbl):
    """Remove all visible borders from a table (IEEE author block style)."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = _mk("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    tblBdr = _mk("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = _mk(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBdr.append(b)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblPr.append(tblBdr)


def _set_table_cell_margins(tbl, top=36, bottom=36, left=115, right=115):
    """
    Set table-wide cell margins in twips.
    Word default is 108 twips top/bottom (~5.4pt per side = 10.8pt per row of dead space).
    36 twips = 1.8pt per side, saving ~14pt per two-row table.
    """
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = _mk("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(existing)
    tblCellMar = _mk("w:tblCellMar")
    for side, w in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        m = _mk(f"w:{side}")
        m.set(qn("w:w"), str(w))
        m.set(qn("w:type"), "dxa")
        tblCellMar.append(m)
    tblPr.append(tblCellMar)


def _fill_author_cell(cell, ordinal: str, name: str, affiliation: str, email: str):
    """Fill one cell of the IEEE author table: ordinal+name bold, affil italic, email italic."""
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    def cp(text: str = "", bold: bool = False, italic: bool = False,
           pt: float = 9, before_pt: float = 0, after_pt: float = 0):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before_pt=before_pt, after_pt=after_pt)
        if text:
            rv = p.add_run(text)
            rv.bold = bold
            rv.italic = italic
            rv.font.name = "Times New Roman"
            rv.font.size = Pt(pt)
        return p

    cp(f"{ordinal} {name}", bold=True, pt=10, before_pt=1, after_pt=0)
    for line in affiliation.split("\n"):
        if line.strip():
            cp(line.strip(), italic=True, pt=9)
    if email:
        cp(email, italic=True, pt=9, after_pt=1)


def _add_author_table(doc: Document, authors: List[Author]):
    """
    Render authors as a borderless table per IEEE two-column format.
    Up to 3 authors per row: 1-3 authors = 1 row, 4-6 = 2 rows, etc.
    Using a table (not sequential paragraphs) preserves correct reading order
    when the DOCX is extracted linearly (e.g., by screen readers or parsers).
    """
    n = len(authors)
    cols = min(3, n)
    rows = math.ceil(n / cols)

    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders_on_table(tbl)
    _set_table_cell_margins(tbl, top=36, bottom=36, left=115, right=115)

    for idx, author in enumerate(authors):
        ri, ci = divmod(idx, cols)
        cell = tbl.cell(ri, ci)
        _fill_author_cell(cell, _ordinal(idx + 1), author.name, author.affiliation, author.email)


# ---------------------------------------------------------------------------
# Section break: single-column → two-column
# ---------------------------------------------------------------------------

def _add_horz_rule_and_section_break(doc: Document):
    """
    Draw a horizontal rule and end the single-column section in the same paragraph.
    Embedding the sectPr in the rule paragraph avoids an extra ~12pt empty paragraph
    that would appear as dead space between the author block and the abstract.
    The document body sectPr is then set to two columns for all content that follows.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(p, before_pt=2, after_pt=0)

    # Bottom border = horizontal rule
    pPr = p._p.get_or_add_pPr()
    pBdr = _mk("w:pBdr")
    b = _mk("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), "000000")
    pBdr.append(b)
    pPr.append(pBdr)

    # Embedded sectPr: marks END of single-column section (continuous break)
    sectPr = _mk("w:sectPr")
    t = _mk("w:type")
    t.set(qn("w:val"), "continuous")
    sectPr.append(t)
    cols1 = _mk("w:cols")
    cols1.set(qn("w:num"), "1")
    cols1.set(qn("w:space"), "720")
    cols1.set(qn("w:equalWidth"), "1")
    sectPr.append(cols1)
    pPr.append(sectPr)

    # Document body sectPr: defines the TWO-COLUMN section for everything after the break
    body_sectPr = doc.element.body.find(qn("w:sectPr"))
    if body_sectPr is not None:
        for c in body_sectPr.findall(qn("w:cols")):
            body_sectPr.remove(c)
        cols2 = _mk("w:cols")
        cols2.set(qn("w:num"), "2")
        cols2.set(qn("w:space"), "360")   # 0.25" column gap
        cols2.set(qn("w:equalWidth"), "1")
        body_sectPr.append(cols2)

    return p


# ---------------------------------------------------------------------------
# Content formatters
# ---------------------------------------------------------------------------

def _add_title(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before_pt=0, after_pt=3)
    run = p.add_run(title)
    _set_font(run, "Times New Roman", 24, bold=True)


def _add_abstract_block(doc: Document, abstract: str, keywords: List[str]):
    """
    Render Abstract and Keywords in 9pt IEEE style.
    Must be called AFTER _add_horz_rule_and_section_break so it lands in the two-column section.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_para_spacing(p, before_pt=8, after_pt=4)

    label = p.add_run("Abstract")
    _set_font(label, "Times New Roman", 9, bold=True)
    body_run = p.add_run(u"—" + abstract)
    _set_font(body_run, "Times New Roman", 9)

    kw_text = ", ".join(keywords)
    pk = doc.add_paragraph()
    pk.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_para_spacing(pk, before_pt=0, after_pt=8)

    kw_label = pk.add_run("Index Terms")
    _set_font(kw_label, "Times New Roman", 9, bold=True, italic=True)
    kw_body = pk.add_run(u"—" + kw_text + ".")
    _set_font(kw_body, "Times New Roman", 9, italic=True)


def _add_section_heading(doc: Document, text: str, level: int, counter: int):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(p, before_pt=8, after_pt=4)

    if level == 1:
        label = f"{_to_roman(counter)}. {text.upper()}"
        run = p.add_run(label)
        _set_font(run, "Times New Roman", 10, bold=True)
    else:
        label = f"{_to_letter(counter)}. {text}"
        run = p.add_run(label)
        _set_font(run, "Times New Roman", 10, italic=True)


def _add_body_paragraph(doc: Document, text: str, first_indent: bool = True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_para_spacing(p, before_pt=0, after_pt=0)

    if first_indent:
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLine"), str(int(0.15 * 1440)))
        pPr.append(ind)

    run = p.add_run(text.strip())
    _set_font(run, "Times New Roman", 10)


def _add_references(doc: Document, references: List[str]):
    heading_p = doc.add_paragraph()
    heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(heading_p, before_pt=8, after_pt=4)
    run = heading_p.add_run("REFERENCES")
    _set_font(run, "Times New Roman", 10, bold=True)

    for idx, ref in enumerate(references, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_para_spacing(p, before_pt=1, after_pt=1)

        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(int(0.25 * 1440)))
        ind.set(qn("w:hanging"), str(int(0.25 * 1440)))
        pPr.append(ind)

        num_run = p.add_run(f"[{idx}] ")
        _set_font(num_run, "Times New Roman", 9, bold=True)
        ref_run = p.add_run(ref)
        _set_font(ref_run, "Times New Roman", 9)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_ieee_docx(paper: IEEEPaper) -> bytes:
    """Return the bytes of a complete IEEE two-column .docx for the given paper."""
    doc = Document()

    # Page setup — IEEE letter spec: T=0.75" B=1.0" L=R=0.625"
    sec = doc.sections[0]
    _set_page_size_letter(sec)
    _set_page_margins(sec, top=0.75, bottom=1.0, left=0.625, right=0.625)

    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    # ── SINGLE-COLUMN SECTION ────────────────────────────────────────────────
    # Title block
    _add_title(doc, paper.title)

    # Author block: borderless 2-row × 3-col table (up to 6 authors)
    # Using a table (not sequential paragraphs) keeps reading order correct
    # when the DOCX is parsed linearly (critical for accessible / machine-read output).
    _add_author_table(doc, paper.authors)

    # Horizontal rule + embedded section break (one paragraph instead of two,
    # saving ~12pt of dead space between the author block and the abstract).
    _add_horz_rule_and_section_break(doc)

    # ── TWO-COLUMN SECTION ───────────────────────────────────────────────────
    # Abstract and keywords land in the two-column section per IEEE format.
    _add_abstract_block(doc, paper.abstract, paper.keywords)

    main_section_counter = 0
    sub_section_counter = 0

    for section in paper.sections:
        if section.level == 1:
            main_section_counter += 1
            sub_section_counter = 0
            _add_section_heading(doc, section.heading, level=1, counter=main_section_counter)
        else:
            sub_section_counter += 1
            _add_section_heading(doc, section.heading, level=2, counter=sub_section_counter)

        paragraphs = re.split(r"\n\n+", section.content.strip())
        for i, para_text in enumerate(paragraphs):
            if para_text.strip():
                _add_body_paragraph(doc, para_text, first_indent=(i == 0))

    _add_references(doc, paper.references)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
